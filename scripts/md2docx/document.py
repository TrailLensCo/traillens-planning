# Copyright (c) 2025 TrailLensCo
# All rights reserved.
#
# This file is proprietary and confidential.

"""Document creation and structure management."""

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.text.paragraph import Paragraph

from md2docx.parser import DocumentMetadata
from md2docx.styles import DocumentStyle, PageSize

logger = logging.getLogger(__name__)


class DocumentBuilder:
    """Builder for creating and structuring Word documents."""

    # Page dimensions in inches
    PAGE_SIZES = {
        PageSize.LETTER: (8.5, 11),
        PageSize.A4: (8.27, 11.69),
    }

    def __init__(self, style: DocumentStyle):
        """
        Initialize the document builder.

        Args:
            style: Document style configuration.
        """
        self.style = style
        self.document = Document()
        self._setup_document()

    def _setup_document(self) -> None:
        """Set up document properties and page layout."""
        # Set up page size and margins
        section = self.document.sections[0]

        # Page size
        width, height = self.PAGE_SIZES.get(
            self.style.page.size, self.PAGE_SIZES[PageSize.LETTER]
        )
        section.page_width = Inches(width)
        section.page_height = Inches(height)

        # Margins
        section.top_margin = Inches(self.style.page.margin_top_inches)
        section.bottom_margin = Inches(self.style.page.margin_bottom_inches)
        section.left_margin = Inches(self.style.page.margin_left_inches)
        section.right_margin = Inches(self.style.page.margin_right_inches)

        # Different first page header/footer
        section.different_first_page_header_footer = (
            self.style.page.different_first_page
        )

        # Set up styles
        self._setup_styles()

        logger.debug("Document setup complete")

    def _setup_styles(self) -> None:
        """Set up document styles."""
        styles = self.document.styles

        # Configure heading styles
        for level in range(1, 7):
            heading_style = self.style.get_heading_style(level)
            style_name = f"Heading {level}"

            try:
                style = styles[style_name]
                font = style.font
                font.name = heading_style.font.name
                font.size = Pt(heading_style.font.size_pt)
                font.bold = heading_style.font.bold
                font.italic = heading_style.font.italic

                if heading_style.font.color:
                    font.color.rgb = RGBColor.from_string(heading_style.font.color)

                pf = style.paragraph_format
                pf.space_before = Pt(heading_style.space_before_pt)
                pf.space_after = Pt(heading_style.space_after_pt)
                pf.keep_with_next = heading_style.keep_with_next

            except KeyError:
                logger.warning(f"Style '{style_name}' not found")

        # Configure Normal style
        try:
            normal = styles["Normal"]
            normal.font.name = self.style.body.font.name
            normal.font.size = Pt(self.style.body.font.size_pt)
            normal.paragraph_format.space_after = Pt(self.style.body.space_after_pt)
            normal.paragraph_format.line_spacing = self.style.body.line_spacing
        except KeyError:
            pass

    def add_title_page(self, metadata: DocumentMetadata) -> None:
        """
        Add a title page to the document.

        Args:
            metadata: Document metadata containing title, author, date.
        """
        if not metadata.title:
            return

        # Add title
        title_para = self.document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Inches(2)

        title_run = title_para.add_run(metadata.title)
        title_run.font.name = self.style.title_font.name
        title_run.font.size = Pt(self.style.title_font.size_pt)
        title_run.bold = self.style.title_font.bold
        if self.style.title_font.color:
            title_run.font.color.rgb = RGBColor.from_string(
                self.style.title_font.color
            )

        # Add author
        if metadata.author:
            author_para = self.document.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_para.paragraph_format.space_before = Inches(0.5)

            author_run = author_para.add_run(metadata.author)
            author_run.font.name = self.style.author_font.name
            author_run.font.size = Pt(self.style.author_font.size_pt)
            if self.style.author_font.color:
                author_run.font.color.rgb = RGBColor.from_string(
                    self.style.author_font.color
                )

        # Add date
        if metadata.date:
            date_para = self.document.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.paragraph_format.space_before = Pt(12)

            date_run = date_para.add_run(metadata.date)
            date_run.font.name = self.style.date_font.name
            date_run.font.size = Pt(self.style.date_font.size_pt)
            if self.style.date_font.color:
                date_run.font.color.rgb = RGBColor.from_string(
                    self.style.date_font.color
                )

        # Add abstract if present
        if metadata.abstract:
            abstract_para = self.document.add_paragraph()
            abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            abstract_para.paragraph_format.space_before = Inches(1)
            abstract_para.paragraph_format.left_indent = Inches(1)
            abstract_para.paragraph_format.right_indent = Inches(1)

            abstract_run = abstract_para.add_run(metadata.abstract)
            abstract_run.italic = True
            abstract_run.font.size = Pt(11)

        # Add page break
        self.document.add_page_break()

        logger.debug("Added title page")

    def add_table_of_contents(self) -> None:
        """Add a table of contents using Word field codes."""
        # Add TOC heading
        toc_heading = self.document.add_paragraph("Table of Contents")
        toc_heading.style = self.document.styles["Heading 1"]
        toc_heading.paragraph_format.page_break_before = False

        # Add TOC field
        paragraph = self.document.add_paragraph()

        # Create field for TOC
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        run = paragraph.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instr_text)

        run = paragraph.add_run()
        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_char_separate)

        # Placeholder text
        run = paragraph.add_run()
        run.text = "Right-click and select 'Update Field' to generate TOC"
        run.italic = True
        run.font.color.rgb = RGBColor.from_string("808080")

        run = paragraph.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)

        # Add page break after TOC
        self.document.add_page_break()

        logger.debug("Added table of contents")

    def setup_headers_footers(
        self, title: Optional[str] = None, date: Optional[str] = None
    ) -> None:
        """
        Set up document headers and footers.

        Args:
            title: Document title for header.
            date: Document date for header.
        """
        section = self.document.sections[0]

        # Set up header (not on first page)
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.clear()

        if title:
            title_run = header_para.add_run(title)
            title_run.font.size = Pt(self.style.header_footer.header_font.size_pt)
            if self.style.header_footer.header_font.color:
                title_run.font.color.rgb = RGBColor.from_string(
                    self.style.header_footer.header_font.color
                )

        if date:
            # Add tab to push date to the right
            header_para.add_run("\t\t")
            date_run = header_para.add_run(date)
            date_run.font.size = Pt(self.style.header_footer.header_font.size_pt)
            if self.style.header_footer.header_font.color:
                date_run.font.color.rgb = RGBColor.from_string(
                    self.style.header_footer.header_font.color
                )

        # Set up footer with page numbers
        if self.style.header_footer.show_page_numbers:
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add "Page X of Y" using field codes
            self._add_page_number_field(footer_para)

        logger.debug("Set up headers and footers")

    def _add_page_number_field(self, paragraph: Paragraph) -> None:
        """
        Add page number field to a paragraph.

        Args:
            paragraph: Target paragraph.
        """
        # "Page "
        run = paragraph.add_run("Page ")
        run.font.size = Pt(self.style.header_footer.footer_font.size_pt)
        if self.style.header_footer.footer_font.color:
            run.font.color.rgb = RGBColor.from_string(
                self.style.header_footer.footer_font.color
            )

        # PAGE field
        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char)

        run = paragraph.add_run()
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        run._r.append(instr)

        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char)

        # " of "
        run = paragraph.add_run(" of ")
        run.font.size = Pt(self.style.header_footer.footer_font.size_pt)
        if self.style.header_footer.footer_font.color:
            run.font.color.rgb = RGBColor.from_string(
                self.style.header_footer.footer_font.color
            )

        # NUMPAGES field
        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char)

        run = paragraph.add_run()
        instr = OxmlElement("w:instrText")
        instr.text = "NUMPAGES"
        run._r.append(instr)

        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char)

    def save(self, path: Path) -> None:
        """
        Save the document to a file.

        Args:
            path: Output file path.
        """
        self.document.save(path)
        logger.info(f"Document saved to {path}")

    def get_document(self) -> Document:
        """
        Get the underlying python-docx Document.

        Returns:
            The Document instance.
        """
        return self.document
