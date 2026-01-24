# Copyright (c) 2025 TrailLensCo
# All rights reserved.
#
# This file is proprietary and confidential.

"""Block element handlers for blockquotes and horizontal rules."""

import logging
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from md2docx.parser import ParsedToken
from md2docx.styles import DocumentStyle

logger = logging.getLogger(__name__)


class BlockHandler:
    """Handler for block elements: blockquotes, horizontal rules, etc."""

    def __init__(self, document: Document, style: DocumentStyle):
        """
        Initialize the block handler.

        Args:
            document: The python-docx Document instance.
            style: Document style configuration.
        """
        self.document = document
        self.style = style
        self._blockquote_depth = 0

    def start_blockquote(self) -> None:
        """Start a blockquote context."""
        self._blockquote_depth += 1

    def end_blockquote(self) -> None:
        """End a blockquote context."""
        if self._blockquote_depth > 0:
            self._blockquote_depth -= 1

    def add_blockquote_paragraph(
        self, tokens: list[ParsedToken], text_handler
    ) -> Paragraph:
        """
        Add a blockquote paragraph.

        Args:
            tokens: Content tokens.
            text_handler: TextHandler for inline formatting.

        Returns:
            The created Paragraph.
        """
        paragraph = self.document.add_paragraph()

        # Apply blockquote styling
        self._style_blockquote_paragraph(paragraph)

        # Add content
        text_handler._add_inline_content(paragraph, tokens)

        # Apply italic to all runs
        for run in paragraph.runs:
            if self.style.blockquote.font.italic:
                run.italic = True

        return paragraph

    def _style_blockquote_paragraph(self, paragraph: Paragraph) -> None:
        """
        Apply blockquote styling to a paragraph.

        Args:
            paragraph: Target paragraph.
        """
        pf = paragraph.paragraph_format

        # Indentation based on nesting depth
        base_indent = self.style.blockquote.left_indent_inches
        indent = base_indent * self._blockquote_depth
        pf.left_indent = Inches(indent)

        # Spacing
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)

        # Add left border
        self._add_left_border(
            paragraph,
            self.style.blockquote.border_color,
            self.style.blockquote.border_width_pt,
        )

    def _add_left_border(
        self, paragraph: Paragraph, color: str, width_pt: float
    ) -> None:
        """
        Add a left border to a paragraph.

        Args:
            paragraph: Target paragraph.
            color: Border color (hex without #).
            width_pt: Border width in points.
        """
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")

        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), str(int(width_pt * 8)))
        left.set(qn("w:color"), color)
        left.set(qn("w:space"), str(int(self.style.blockquote.padding_pt)))
        pBdr.append(left)

        pPr.append(pBdr)

    def add_horizontal_rule(self, as_page_break: bool = False) -> Paragraph:
        """
        Add a horizontal rule to the document.

        Args:
            as_page_break: Insert a page break instead of a line.

        Returns:
            The created Paragraph.
        """
        if as_page_break:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)
            return paragraph

        # Create a paragraph with a bottom border to simulate HR
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)

        # Add bottom border
        self._add_bottom_border(paragraph, "CCCCCC", 1.0)

        return paragraph

    def _add_bottom_border(
        self, paragraph: Paragraph, color: str, width_pt: float
    ) -> None:
        """
        Add a bottom border to a paragraph.

        Args:
            paragraph: Target paragraph.
            color: Border color (hex without #).
            width_pt: Border width in points.
        """
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(int(width_pt * 8)))
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)

        pPr.append(pBdr)

    @property
    def in_blockquote(self) -> bool:
        """Check if currently inside a blockquote."""
        return self._blockquote_depth > 0

    @property
    def blockquote_depth(self) -> int:
        """Get the current blockquote nesting depth."""
        return self._blockquote_depth


class FootnoteHandler:
    """Handler for footnotes and endnotes."""

    def __init__(self, document: Document, style: DocumentStyle):
        """
        Initialize the footnote handler.

        Args:
            document: The python-docx Document instance.
            style: Document style configuration.
        """
        self.document = document
        self.style = style
        self._footnotes: dict[str, str] = {}
        self._footnote_counter = 0

    def register_footnote(self, ref: str, content: str) -> None:
        """
        Register a footnote definition.

        Args:
            ref: Footnote reference identifier.
            content: Footnote content text.
        """
        self._footnotes[ref] = content
        logger.debug(f"Registered footnote: {ref}")

    def add_footnote_reference(self, ref: str, paragraph: Paragraph) -> None:
        """
        Add a footnote reference in the text.

        Args:
            ref: Footnote reference identifier.
            paragraph: Target paragraph.
        """
        self._footnote_counter += 1

        # Add superscript reference number
        run = paragraph.add_run(f"[{self._footnote_counter}]")
        run.font.superscript = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("0563C1")

    def get_footnote_content(self, ref: str) -> Optional[str]:
        """
        Get the content of a footnote.

        Args:
            ref: Footnote reference identifier.

        Returns:
            Footnote content or None if not found.
        """
        return self._footnotes.get(ref)

    def add_footnotes_section(self) -> None:
        """Add a footnotes section at the end of the document."""
        if not self._footnotes:
            return

        # Add a horizontal rule
        hr = self.document.add_paragraph()
        hr.paragraph_format.space_before = Pt(24)

        pPr = hr._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "4")
        top.set(qn("w:color"), "CCCCCC")
        pBdr.append(top)
        pPr.append(pBdr)

        # Add footnote entries
        for i, (ref, content) in enumerate(self._footnotes.items(), 1):
            para = self.document.add_paragraph()
            para.paragraph_format.space_after = Pt(4)

            # Add number
            num_run = para.add_run(f"[{i}] ")
            num_run.font.size = Pt(9)
            num_run.bold = True

            # Add content
            content_run = para.add_run(content)
            content_run.font.size = Pt(9)


# Import docx for page break enum
import docx
