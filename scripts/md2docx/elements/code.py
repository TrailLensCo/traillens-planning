# Copyright (c) 2025 TrailLensCo
# All rights reserved.
#
# This file is proprietary and confidential.

"""Code block element handlers with syntax highlighting."""

import logging
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.text.paragraph import Paragraph

from md2docx.styles import DocumentStyle

logger = logging.getLogger(__name__)

# Optional pygments import for syntax highlighting
try:
    from pygments import highlight
    from pygments.formatters import HtmlFormatter
    from pygments.lexers import get_lexer_by_name, guess_lexer
    from pygments.token import Token
    from pygments.util import ClassNotFound

    PYGMENTS_AVAILABLE = True
except ImportError:
    PYGMENTS_AVAILABLE = False
    logger.warning("Pygments not available, syntax highlighting disabled")


# Token type to color mapping
TOKEN_COLORS = {
    "Comment": "6A9955",       # Green
    "String": "CE9178",         # Orange
    "Number": "B5CEA8",         # Light green
    "Keyword": "569CD6",        # Blue
    "Name.Function": "DCDCAA", # Yellow
    "Name.Class": "4EC9B0",    # Cyan
    "Name.Builtin": "4FC1FF",  # Light blue
    "Operator": "D4D4D4",      # Light gray
    "Punctuation": "D4D4D4",   # Light gray
    "Name.Decorator": "DCDCAA", # Yellow
    "Name.Exception": "4EC9B0", # Cyan
}


class CodeHandler:
    """Handler for code blocks with syntax highlighting."""

    def __init__(self, document: Document, style: DocumentStyle):
        """
        Initialize the code handler.

        Args:
            document: The python-docx Document instance.
            style: Document style configuration.
        """
        self.document = document
        self.style = style

    def add_code_block(
        self,
        code: str,
        language: Optional[str] = None,
        highlight_syntax: bool = True,
    ) -> list[Paragraph]:
        """
        Add a code block to the document.

        Args:
            code: Code content.
            language: Language identifier for syntax highlighting.
            highlight_syntax: Enable syntax highlighting.

        Returns:
            List of created Paragraphs (one per line).
        """
        paragraphs = []

        # Split code into lines, preserving empty lines
        lines = code.split("\n")

        # Remove trailing empty line if present (common in fenced blocks)
        if lines and lines[-1] == "":
            lines = lines[:-1]

        # Get syntax tokens if highlighting enabled
        tokens_by_line = None
        if highlight_syntax and PYGMENTS_AVAILABLE and language:
            tokens_by_line = self._get_syntax_tokens(code, language)

        for line_num, line in enumerate(lines):
            paragraph = self.document.add_paragraph()
            self._style_code_paragraph(paragraph, is_first=(line_num == 0),
                                       is_last=(line_num == len(lines) - 1))

            if tokens_by_line and line_num < len(tokens_by_line):
                # Add highlighted tokens
                for token_type, token_text in tokens_by_line[line_num]:
                    self._add_highlighted_run(paragraph, token_text, token_type)
            else:
                # Add plain text
                if line:
                    run = paragraph.add_run(line)
                    self._style_code_run(run)
                else:
                    # Empty line - add non-breaking space to preserve line
                    run = paragraph.add_run("\u00A0")
                    self._style_code_run(run)

            paragraphs.append(paragraph)

        logger.debug(f"Added code block with {len(lines)} lines")
        return paragraphs

    def add_inline_code(self, text: str, paragraph: Paragraph) -> None:
        """
        Add inline code to a paragraph.

        Args:
            text: Code text.
            paragraph: Target paragraph.
        """
        run = paragraph.add_run(text)
        run.font.name = self.style.code.font.name
        run.font.size = Pt(self.style.code.font.size_pt)

        # Add background shading
        self._add_run_background(run, "F0F0F0")

    def _style_code_paragraph(
        self, paragraph: Paragraph, is_first: bool = False, is_last: bool = False
    ) -> None:
        """
        Apply code block styling to a paragraph.

        Args:
            paragraph: Target paragraph.
            is_first: True if this is the first line.
            is_last: True if this is the last line.
        """
        pf = paragraph.paragraph_format

        # Remove spacing between code lines
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

        # Add padding on first/last lines
        if is_first:
            pf.space_before = Pt(self.style.code.padding_pt)
        if is_last:
            pf.space_after = Pt(self.style.code.padding_pt)

        # Add left indent for padding
        pf.left_indent = Pt(self.style.code.padding_pt * 2)
        pf.right_indent = Pt(self.style.code.padding_pt * 2)

        # Keep lines together
        pf.keep_together = True

        # Add background shading
        self._add_paragraph_shading(paragraph, self.style.code.background_color)

        # Add border
        self._add_paragraph_border(
            paragraph,
            self.style.code.border_color,
            self.style.code.border_width_pt,
            is_first,
            is_last,
        )

    def _style_code_run(self, run) -> None:
        """
        Apply code font styling to a run.

        Args:
            run: Target run.
        """
        run.font.name = self.style.code.font.name
        run.font.size = Pt(self.style.code.font.size_pt)

    def _add_highlighted_run(
        self, paragraph: Paragraph, text: str, token_type
    ) -> None:
        """
        Add a syntax-highlighted run to a paragraph.

        Args:
            paragraph: Target paragraph.
            text: Token text.
            token_type: Pygments token type.
        """
        run = paragraph.add_run(text)
        self._style_code_run(run)

        # Get color for token type
        color = self._get_token_color(token_type)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    def _get_token_color(self, token_type) -> Optional[str]:
        """
        Get the color for a Pygments token type.

        Args:
            token_type: Pygments token type.

        Returns:
            Hex color string or None.
        """
        if not PYGMENTS_AVAILABLE:
            return None

        # Convert token type to string representation
        token_str = str(token_type)

        # Check for direct match
        for key, color in TOKEN_COLORS.items():
            if key in token_str:
                return color

        # Check token hierarchy
        if Token.Comment in token_type:
            return TOKEN_COLORS["Comment"]
        elif Token.String in token_type:
            return TOKEN_COLORS["String"]
        elif Token.Number in token_type:
            return TOKEN_COLORS["Number"]
        elif Token.Keyword in token_type:
            return TOKEN_COLORS["Keyword"]
        elif Token.Name.Function in token_type:
            return TOKEN_COLORS["Name.Function"]
        elif Token.Name.Class in token_type:
            return TOKEN_COLORS["Name.Class"]
        elif Token.Name.Builtin in token_type:
            return TOKEN_COLORS["Name.Builtin"]

        return None

    def _get_syntax_tokens(
        self, code: str, language: str
    ) -> list[list[tuple]]:
        """
        Get syntax tokens organized by line.

        Args:
            code: Source code.
            language: Language identifier.

        Returns:
            List of lists, where each inner list contains (token_type, text) tuples.
        """
        if not PYGMENTS_AVAILABLE:
            return []

        try:
            lexer = get_lexer_by_name(language)
        except ClassNotFound:
            try:
                lexer = guess_lexer(code)
            except ClassNotFound:
                return []

        tokens = list(lexer.get_tokens(code))

        # Organize by line
        lines = [[]]
        for token_type, token_text in tokens:
            if "\n" in token_text:
                parts = token_text.split("\n")
                for i, part in enumerate(parts):
                    if i > 0:
                        lines.append([])
                    if part:
                        lines[-1].append((token_type, part))
            else:
                lines[-1].append((token_type, token_text))

        return lines

    def _add_paragraph_shading(self, paragraph: Paragraph, color: str) -> None:
        """
        Add background shading to a paragraph.

        Args:
            paragraph: Target paragraph.
            color: Background color (hex without #).
        """
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        pPr.append(shd)

    def _add_paragraph_border(
        self,
        paragraph: Paragraph,
        color: str,
        width_pt: float,
        is_first: bool,
        is_last: bool,
    ) -> None:
        """
        Add borders to a paragraph.

        Args:
            paragraph: Target paragraph.
            color: Border color (hex without #).
            width_pt: Border width in points.
            is_first: True if first line of block.
            is_last: True if last line of block.
        """
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")

        size = str(int(width_pt * 8))  # Convert to eighths of a point

        # Always add left and right borders
        for side in ["left", "right"]:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), size)
            border.set(qn("w:color"), color)
            border.set(qn("w:space"), "4")
            pBdr.append(border)

        # Add top border only on first line
        if is_first:
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), size)
            top.set(qn("w:color"), color)
            pBdr.append(top)

        # Add bottom border only on last line
        if is_last:
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), size)
            bottom.set(qn("w:color"), color)
            pBdr.append(bottom)

        pPr.append(pBdr)

    def _add_run_background(self, run, color: str) -> None:
        """
        Add background shading to a run.

        Args:
            run: Target run.
            color: Background color (hex without #).
        """
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        run._r.get_or_add_rPr().append(shd)
