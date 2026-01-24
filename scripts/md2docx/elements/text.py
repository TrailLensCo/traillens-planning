# Copyright (c) 2025 TrailLensCo
# All rights reserved.
#
# This file is proprietary and confidential.

"""Text element handlers for paragraphs, headings, and inline formatting."""

import logging
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from md2docx.parser import ParsedToken, extract_text_content
from md2docx.styles import DocumentStyle, FontStyle, HeadingStyle, TextAlignment

logger = logging.getLogger(__name__)


class TextHandler:
    """Handler for text elements: paragraphs, headings, and inline formatting."""

    def __init__(self, document: Document, style: DocumentStyle):
        """
        Initialize the text handler.

        Args:
            document: The python-docx Document instance.
            style: Document style configuration.
        """
        self.document = document
        self.style = style

    def add_heading(self, level: int, tokens: list[ParsedToken]) -> Paragraph:
        """
        Add a heading to the document.

        Args:
            level: Heading level (1-6).
            tokens: Child tokens containing heading content.

        Returns:
            The created Paragraph.
        """
        level = max(1, min(6, level))  # Clamp to valid range
        heading_style = self.style.get_heading_style(level)

        # Create paragraph with Word's built-in heading style
        paragraph = self.document.add_paragraph(style=f"Heading {level}")

        # Apply custom styling
        self._apply_heading_style(paragraph, heading_style)

        # Add content with formatting
        self._add_inline_content(paragraph, tokens)

        logger.debug(f"Added heading level {level}")
        return paragraph

    def add_paragraph(self, tokens: list[ParsedToken]) -> Paragraph:
        """
        Add a regular paragraph to the document.

        Args:
            tokens: Child tokens containing paragraph content.

        Returns:
            The created Paragraph.
        """
        paragraph = self.document.add_paragraph()

        # Apply body text styling
        self._apply_paragraph_style(paragraph)

        # Add content with formatting
        self._add_inline_content(paragraph, tokens)

        return paragraph

    def add_text_run(
        self,
        paragraph: Paragraph,
        text: str,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        strikethrough: bool = False,
        code: bool = False,
        color: Optional[str] = None,
    ) -> Run:
        """
        Add a text run with formatting to a paragraph.

        Args:
            paragraph: Target paragraph.
            text: Text content.
            bold: Apply bold formatting.
            italic: Apply italic formatting.
            underline: Apply underline formatting.
            strikethrough: Apply strikethrough formatting.
            code: Apply inline code formatting.
            color: Text color (hex without #).

        Returns:
            The created Run.
        """
        run = paragraph.add_run(text)

        run.bold = bold
        run.italic = italic
        run.underline = underline

        if strikethrough:
            run.font.strike = True

        if code:
            run.font.name = self.style.code.font.name
            run.font.size = Pt(self.style.code.font.size_pt)
            # Add light background for inline code
            self._add_run_shading(run, "F0F0F0")

        if color:
            run.font.color.rgb = RGBColor.from_string(color)

        return run

    def _add_inline_content(
        self,
        paragraph: Paragraph,
        tokens: list[ParsedToken],
        format_state: Optional[dict] = None,
    ) -> None:
        """
        Add inline content with formatting to a paragraph.

        Args:
            paragraph: Target paragraph.
            tokens: Tokens to process.
            format_state: Current formatting state (bold, italic, etc.).
        """
        if format_state is None:
            format_state = {
                "bold": False,
                "italic": False,
                "strikethrough": False,
                "code": False,
                "link_href": None,
            }

        for token in tokens:
            if token.type == "text":
                run = self.add_text_run(
                    paragraph,
                    token.content,
                    bold=format_state["bold"],
                    italic=format_state["italic"],
                    strikethrough=format_state["strikethrough"],
                    code=format_state["code"],
                )
                # Handle link formatting
                if format_state.get("link_href"):
                    self._add_hyperlink(
                        paragraph,
                        format_state["link_href"],
                        run,
                    )

            elif token.type == "code_inline":
                self.add_text_run(
                    paragraph,
                    token.content,
                    code=True,
                    bold=format_state["bold"],
                    italic=format_state["italic"],
                )

            elif token.type == "softbreak":
                paragraph.add_run(" ")

            elif token.type == "hardbreak":
                run = paragraph.add_run()
                run.add_break()

            elif token.type == "strong_open":
                format_state["bold"] = True

            elif token.type == "strong_close":
                format_state["bold"] = False

            elif token.type == "em_open":
                format_state["italic"] = True

            elif token.type == "em_close":
                format_state["italic"] = False

            elif token.type == "s_open":
                format_state["strikethrough"] = True

            elif token.type == "s_close":
                format_state["strikethrough"] = False

            elif token.type == "link_open":
                format_state["link_href"] = token.attrs.get("href", "")

            elif token.type == "link_close":
                format_state["link_href"] = None

            elif token.type == "image":
                # Images handled separately by MediaHandler
                pass

            elif token.children:
                # Process nested tokens
                self._add_inline_content(paragraph, token.children, format_state)

    def _apply_heading_style(
        self, paragraph: Paragraph, heading_style: HeadingStyle
    ) -> None:
        """
        Apply custom heading style to a paragraph.

        Args:
            paragraph: Target paragraph.
            heading_style: Heading style configuration.
        """
        pf = paragraph.paragraph_format

        # Spacing
        pf.space_before = Pt(heading_style.space_before_pt)
        pf.space_after = Pt(heading_style.space_after_pt)
        pf.keep_with_next = heading_style.keep_with_next

        # Apply font styling to any existing runs
        for run in paragraph.runs:
            self._apply_font_style(run, heading_style.font)

    def _apply_paragraph_style(self, paragraph: Paragraph) -> None:
        """
        Apply body paragraph styling.

        Args:
            paragraph: Target paragraph.
        """
        pf = paragraph.paragraph_format

        # Spacing
        pf.space_before = Pt(self.style.body.space_before_pt)
        pf.space_after = Pt(self.style.body.space_after_pt)

        # Line spacing
        pf.line_spacing = self.style.body.line_spacing

        # Alignment
        alignment_map = {
            TextAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
            TextAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
            TextAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
            TextAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        pf.alignment = alignment_map.get(
            self.style.body_alignment, WD_ALIGN_PARAGRAPH.LEFT
        )

    def _apply_font_style(self, run: Run, font_style: FontStyle) -> None:
        """
        Apply font styling to a run.

        Args:
            run: Target run.
            font_style: Font style configuration.
        """
        run.font.name = font_style.name
        run.font.size = Pt(font_style.size_pt)
        run.bold = font_style.bold
        run.italic = font_style.italic
        run.underline = font_style.underline

        if font_style.strikethrough:
            run.font.strike = True

        if font_style.color:
            run.font.color.rgb = RGBColor.from_string(font_style.color)

    def _add_hyperlink(
        self, paragraph: Paragraph, url: str, run: Run
    ) -> None:
        """
        Convert a run to a hyperlink.

        Args:
            paragraph: Parent paragraph.
            url: Link URL.
            run: Run to make into a hyperlink.
        """
        # Apply hyperlink color
        run.font.color.rgb = RGBColor.from_string(self.style.hyperlink_color)

        # Create hyperlink element
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        # Create hyperlink XML element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Move the run into the hyperlink
        run._r.getparent().replace(run._r, hyperlink)
        hyperlink.append(run._r)

    def _add_run_shading(self, run: Run, color: str) -> None:
        """
        Add background shading to a run.

        Args:
            run: Target run.
            color: Background color (hex without #).
        """
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        run._r.get_or_add_rPr().append(shd)
