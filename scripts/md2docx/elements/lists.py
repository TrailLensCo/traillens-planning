# Copyright (c) 2025 TrailLensCo
# All rights reserved.
#
# This file is proprietary and confidential.

"""List element handlers for bullet and numbered lists."""

import logging
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from md2docx.parser import ParsedToken, extract_text_content, parse_task_list_item
from md2docx.styles import DocumentStyle

logger = logging.getLogger(__name__)


class ListHandler:
    """Handler for list elements: bullet lists, numbered lists, task lists."""

    # Checkbox characters
    CHECKBOX_UNCHECKED = "☐"
    CHECKBOX_CHECKED = "☑"

    def __init__(self, document: Document, style: DocumentStyle):
        """
        Initialize the list handler.

        Args:
            document: The python-docx Document instance.
            style: Document style configuration.
        """
        self.document = document
        self.style = style
        self._list_stack: list[dict] = []
        self._num_id_counter = 1

    def start_list(self, ordered: bool = False) -> None:
        """
        Start a new list context.

        Args:
            ordered: True for numbered list, False for bullet list.
        """
        level = len(self._list_stack)
        self._list_stack.append({
            "ordered": ordered,
            "level": level,
            "item_count": 0,
        })
        logger.debug(f"Started {'ordered' if ordered else 'bullet'} list at level {level}")

    def end_list(self) -> None:
        """End the current list context."""
        if self._list_stack:
            self._list_stack.pop()
            logger.debug(f"Ended list, depth now {len(self._list_stack)}")

    def add_list_item(
        self,
        tokens: list[ParsedToken],
        text_handler,
        is_task: bool = False,
        checked: bool = False,
    ) -> Paragraph:
        """
        Add a list item.

        Args:
            tokens: Content tokens for the list item.
            text_handler: TextHandler instance for inline formatting.
            is_task: True if this is a task list item.
            checked: True if task checkbox is checked.

        Returns:
            The created Paragraph.
        """
        if not self._list_stack:
            # Not in a list context, treat as regular paragraph
            return text_handler.add_paragraph(tokens)

        list_info = self._list_stack[-1]
        list_info["item_count"] += 1

        level = min(list_info["level"], 2)  # Max 3 levels (0, 1, 2)
        ordered = list_info["ordered"]

        # Create paragraph with appropriate list style
        if ordered:
            style_name = f"List Number{' ' + str(level + 1) if level > 0 else ''}"
        else:
            style_name = f"List Bullet{' ' + str(level + 1) if level > 0 else ''}"

        try:
            paragraph = self.document.add_paragraph(style=style_name)
        except KeyError:
            # Style doesn't exist, create manually
            paragraph = self.document.add_paragraph()
            self._apply_list_formatting(paragraph, ordered, level)

        # Apply indentation for nested lists
        self._apply_indent(paragraph, level)

        # Handle task list items
        if is_task:
            checkbox = self.CHECKBOX_CHECKED if checked else self.CHECKBOX_UNCHECKED
            run = paragraph.add_run(f"{checkbox} ")
            run.font.name = "Segoe UI Symbol"  # Better checkbox rendering
            # Get content without checkbox
            _, content = parse_task_list_item(tokens[0] if tokens else ParsedToken(type=""))
            text_handler.add_text_run(paragraph, content)
        else:
            # Add content with inline formatting
            text_handler._add_inline_content(paragraph, self._get_inline_tokens(tokens))

        return paragraph

    def _get_inline_tokens(self, tokens: list[ParsedToken]) -> list[ParsedToken]:
        """
        Extract inline tokens from list item tokens.

        Args:
            tokens: List item tokens.

        Returns:
            List of inline tokens for content rendering.
        """
        inline_tokens = []
        for token in tokens:
            if token.type == "inline":
                inline_tokens.extend(token.children if token.children else [token])
            elif token.type == "paragraph_open":
                continue
            elif token.type == "paragraph_close":
                continue
            elif token.children:
                inline_tokens.extend(self._get_inline_tokens(token.children))
        return inline_tokens

    def _apply_list_formatting(
        self, paragraph: Paragraph, ordered: bool, level: int
    ) -> None:
        """
        Apply list formatting to a paragraph.

        Args:
            paragraph: Target paragraph.
            ordered: True for numbered list.
            level: Nesting level (0-based).
        """
        # Set left indent based on level
        indent = 0.25 + (level * 0.25)  # inches
        paragraph.paragraph_format.left_indent = Inches(indent)

        # Add numbering properties
        self._set_numbering(paragraph, ordered, level)

    def _apply_indent(self, paragraph: Paragraph, level: int) -> None:
        """
        Apply indentation for list level.

        Args:
            paragraph: Target paragraph.
            level: Nesting level (0-based).
        """
        indent = 0.25 + (level * 0.25)  # inches
        paragraph.paragraph_format.left_indent = Inches(indent)

    def _set_numbering(
        self, paragraph: Paragraph, ordered: bool, level: int
    ) -> None:
        """
        Set numbering/bullet properties on a paragraph.

        Args:
            paragraph: Target paragraph.
            ordered: True for numbered list.
            level: Nesting level.
        """
        # Create numbering properties
        pPr = paragraph._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")

        # Set indent level
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        numPr.append(ilvl)

        # Set numbering ID (use abstract num for bullet vs numbered)
        numId = OxmlElement("w:numId")
        num_val = "1" if not ordered else "2"
        numId.set(qn("w:val"), num_val)
        numPr.append(numId)

        pPr.append(numPr)

    @property
    def current_level(self) -> int:
        """Get the current list nesting level."""
        return len(self._list_stack) - 1 if self._list_stack else -1

    @property
    def in_list(self) -> bool:
        """Check if currently inside a list."""
        return len(self._list_stack) > 0


class DefinitionListHandler:
    """Handler for definition lists (term/definition pairs)."""

    def __init__(self, document: Document, style: DocumentStyle):
        """
        Initialize the definition list handler.

        Args:
            document: The python-docx Document instance.
            style: Document style configuration.
        """
        self.document = document
        self.style = style

    def add_term(self, text: str, text_handler) -> Paragraph:
        """
        Add a definition term.

        Args:
            text: Term text.
            text_handler: TextHandler for formatting.

        Returns:
            The created Paragraph.
        """
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        paragraph.paragraph_format.space_after = Pt(2)
        return paragraph

    def add_definition(self, tokens: list[ParsedToken], text_handler) -> Paragraph:
        """
        Add a definition.

        Args:
            tokens: Definition content tokens.
            text_handler: TextHandler for formatting.

        Returns:
            The created Paragraph.
        """
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.space_after = Pt(6)
        text_handler._add_inline_content(paragraph, tokens)
        return paragraph
