# Copyright (c) 2025 TrailLensCo
# All rights reserved.
#
# This file is proprietary and confidential.

"""Media element handlers for images and diagrams."""

import io
import logging
from pathlib import Path
from typing import Optional, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image

from md2docx.styles import DocumentStyle
from md2docx.utils import download_file, is_url, resolve_path, TempFileManager

logger = logging.getLogger(__name__)


class MediaHandler:
    """Handler for media elements: images and diagrams."""

    def __init__(
        self,
        document: Document,
        style: DocumentStyle,
        base_dir: Optional[Path] = None,
        temp_manager: Optional[TempFileManager] = None,
    ):
        """
        Initialize the media handler.

        Args:
            document: The python-docx Document instance.
            style: Document style configuration.
            base_dir: Base directory for resolving relative paths.
            temp_manager: Temporary file manager.
        """
        self.document = document
        self.style = style
        self.base_dir = base_dir or Path.cwd()
        self.temp_manager = temp_manager or TempFileManager()

    def add_image(
        self,
        src: str,
        alt: str = "",
        title: str = "",
        max_width: Optional[float] = None,
    ) -> Optional[Paragraph]:
        """
        Add an image to the document.

        Args:
            src: Image source (file path or URL).
            alt: Alt text for the image.
            title: Image title.
            max_width: Maximum width in inches.

        Returns:
            The paragraph containing the image, or None if failed.
        """
        max_width = max_width or self.style.max_image_width_inches

        # Resolve image path
        image_path = self._resolve_image(src)
        if image_path is None:
            logger.warning(f"Could not resolve image: {src}")
            return self._add_image_placeholder(src, alt)

        try:
            # Get image dimensions
            width, height = self._get_scaled_dimensions(image_path, max_width)

            # Create centered paragraph for image
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add image
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(width))

            logger.debug(f"Added image: {src}")

            # Add caption if alt text provided
            if alt:
                self._add_caption(alt)

            return paragraph

        except Exception as e:
            logger.error(f"Failed to add image {src}: {e}")
            return self._add_image_placeholder(src, alt)

    def add_diagram_image(
        self,
        image_path: Path,
        caption: str = "",
        max_width: Optional[float] = None,
    ) -> Optional[Paragraph]:
        """
        Add a rendered diagram image to the document.

        Args:
            image_path: Path to the rendered diagram image.
            caption: Caption for the diagram.
            max_width: Maximum width in inches.

        Returns:
            The paragraph containing the image, or None if failed.
        """
        return self.add_image(str(image_path), alt=caption, max_width=max_width)

    def _resolve_image(self, src: str) -> Optional[Path]:
        """
        Resolve an image source to a local file path.

        Args:
            src: Image source (path or URL).

        Returns:
            Path to the image file, or None if unresolvable.
        """
        if is_url(src):
            # Download image to temp file
            suffix = Path(src).suffix or ".png"
            temp_path = self.temp_manager.create_temp_file(suffix=suffix)
            return download_file(src, temp_path)
        else:
            # Resolve local path
            return resolve_path(src, self.base_dir)

    def _get_scaled_dimensions(
        self, image_path: Path, max_width: float
    ) -> tuple[float, float]:
        """
        Get scaled dimensions for an image.

        Args:
            image_path: Path to the image.
            max_width: Maximum width in inches.

        Returns:
            Tuple of (width, height) in inches.
        """
        try:
            with Image.open(image_path) as img:
                # Get dimensions in inches (assuming 96 DPI if not specified)
                dpi = img.info.get("dpi", (96, 96))
                if isinstance(dpi, tuple):
                    dpi_x, dpi_y = dpi
                else:
                    dpi_x = dpi_y = dpi

                width_inches = img.width / dpi_x
                height_inches = img.height / dpi_y

                # Scale if necessary
                if width_inches > max_width:
                    scale = max_width / width_inches
                    width_inches = max_width
                    height_inches *= scale

                return width_inches, height_inches

        except Exception as e:
            logger.warning(f"Could not get image dimensions: {e}")
            # Return default dimensions
            return max_width, max_width * 0.75

    def _add_caption(self, caption: str) -> Paragraph:
        """
        Add an image caption.

        Args:
            caption: Caption text.

        Returns:
            The caption paragraph.
        """
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(12)

        run = paragraph.add_run(caption)
        run.font.name = self.style.image_caption_font.name
        run.font.size = Pt(self.style.image_caption_font.size_pt)
        run.italic = self.style.image_caption_font.italic

        if self.style.image_caption_font.color:
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor.from_string(
                self.style.image_caption_font.color
            )

        return paragraph

    def _add_image_placeholder(
        self, src: str, alt: str = ""
    ) -> Paragraph:
        """
        Add a placeholder for a missing image.

        Args:
            src: Original image source.
            alt: Alt text.

        Returns:
            The placeholder paragraph.
        """
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

        # Add placeholder text in a box
        run = paragraph.add_run(f"[Image: {alt or src}]")
        run.italic = True
        run.font.color.rgb = RGBColor.from_string("808080")

        return paragraph


# Import RGBColor for placeholder
from docx.shared import RGBColor
