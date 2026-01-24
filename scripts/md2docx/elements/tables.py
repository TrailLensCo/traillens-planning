# Copyright (c) 2025 TrailLensCo
# All rights reserved.
#
# This file is proprietary and confidential.

"""Table element handlers."""

import logging
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.table import Table, _Cell

from md2docx.parser import ParsedToken, extract_text_content
from md2docx.styles import DocumentStyle

logger = logging.getLogger(__name__)


class TableHandler:
    """Handler for table elements."""

    def __init__(self, document: Document, style: DocumentStyle):
        """
        Initialize the table handler.

        Args:
            document: The python-docx Document instance.
            style: Document style configuration.
        """
        self.document = document
        self.style = style
        self._current_table: Optional[Table] = None
        self._current_row: int = 0
        self._column_alignments: list[str] = []
        self._in_header: bool = False

    def create_table(self, rows: int, cols: int) -> Table:
        """
        Create a new table.

        Args:
            rows: Number of rows.
            cols: Number of columns.

        Returns:
            The created Table.
        """
        self._current_table = self.document.add_table(rows=rows, cols=cols)
        self._current_row = 0
        self._column_alignments = [""] * cols

        # Apply table styling
        self._apply_table_style()

        logger.debug(f"Created table with {rows} rows and {cols} columns")
        return self._current_table

    def start_table(self) -> None:
        """Start building a table (for streaming parsing)."""
        self._current_table = None
        self._current_row = 0
        self._column_alignments = []
        self._in_header = False

    def end_table(self) -> Optional[Table]:
        """
        End table building.

        Returns:
            The completed table.
        """
        table = self._current_table
        self._current_table = None
        self._current_row = 0
        return table

    def start_header(self) -> None:
        """Mark the start of header rows."""
        self._in_header = True

    def end_header(self) -> None:
        """Mark the end of header rows."""
        self._in_header = False

    def set_alignments(self, alignments: list[str]) -> None:
        """
        Set column alignments.

        Args:
            alignments: List of alignment strings ('left', 'center', 'right').
        """
        self._column_alignments = alignments

    def add_row(
        self, cells_content: list[list[ParsedToken]], text_handler
    ) -> None:
        """
        Add a row to the current table.

        Args:
            cells_content: List of token lists, one per cell.
            text_handler: TextHandler for formatting cell content.
        """
        if self._current_table is None:
            # Create table on first row
            cols = len(cells_content)
            self._current_table = self.document.add_table(rows=0, cols=cols)
            self._column_alignments = [""] * cols
            self._apply_table_style()

        # Add a new row
        row = self._current_table.add_row()
        self._current_row += 1

        # Fill cells
        for col_idx, cell_tokens in enumerate(cells_content):
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                self._fill_cell(
                    cell,
                    cell_tokens,
                    text_handler,
                    col_idx,
                    is_header=self._in_header,
                )

        # Apply row styling
        if self._in_header:
            self._style_header_row(row)
        else:
            # Zebra striping for body rows
            row_num = self._current_row - 1  # Adjust for header
            if row_num > 0 and row_num % 2 == 0:
                self._apply_row_shading(row, self.style.table.row_alt_background)

    def _fill_cell(
        self,
        cell: _Cell,
        tokens: list[ParsedToken],
        text_handler,
        col_idx: int,
        is_header: bool = False,
    ) -> None:
        """
        Fill a table cell with content.

        Args:
            cell: Target cell.
            tokens: Content tokens.
            text_handler: TextHandler for formatting.
            col_idx: Column index for alignment.
            is_header: True if this is a header cell.
        """
        # Clear any default paragraph
        cell._tc.clear_content()

        # Add paragraph with content
        paragraph = cell.add_paragraph()

        # Apply alignment
        alignment = (
            self._column_alignments[col_idx]
            if col_idx < len(self._column_alignments)
            else ""
        )
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        paragraph.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # Add content
        inline_tokens = self._extract_inline(tokens)
        text_handler._add_inline_content(paragraph, inline_tokens)

        # Apply header formatting
        if is_header:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(
                    self.style.table.header_text_color
                )

        # Apply cell padding
        self._apply_cell_padding(cell)

    def _extract_inline(self, tokens: list[ParsedToken]) -> list[ParsedToken]:
        """
        Extract inline tokens from cell content.

        Args:
            tokens: Cell tokens.

        Returns:
            Inline tokens for rendering.
        """
        inline_tokens = []
        for token in tokens:
            if token.type == "inline":
                inline_tokens.extend(token.children if token.children else [token])
            elif token.children:
                inline_tokens.extend(self._extract_inline(token.children))
        return inline_tokens

    def _apply_table_style(self) -> None:
        """Apply table styling."""
        if not self._current_table:
            return

        table = self._current_table

        # Auto-fit table width
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Apply borders
        self._apply_table_borders()

    def _apply_table_borders(self) -> None:
        """Apply borders to the table."""
        if not self._current_table:
            return

        tbl = self._current_table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        # Create borders element
        tblBorders = OxmlElement("w:tblBorders")

        border_color = self.style.table.border_color
        border_size = str(int(self.style.table.border_width_pt * 8))  # Convert to eighths of a point

        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), border_size)
            border.set(qn("w:color"), border_color)
            tblBorders.append(border)

        tblPr.append(tblBorders)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _style_header_row(self, row) -> None:
        """
        Apply header row styling.

        Args:
            row: Table row to style.
        """
        for cell in row.cells:
            self._set_cell_background(cell, self.style.table.header_background)

    def _apply_row_shading(self, row, color: str) -> None:
        """
        Apply background shading to a row.

        Args:
            row: Table row.
            color: Background color (hex without #).
        """
        for cell in row.cells:
            self._set_cell_background(cell, color)

    def _set_cell_background(self, cell: _Cell, color: str) -> None:
        """
        Set cell background color.

        Args:
            cell: Target cell.
            color: Background color (hex without #).
        """
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

    def _apply_cell_padding(self, cell: _Cell) -> None:
        """
        Apply padding to a cell.

        Args:
            cell: Target cell.
        """
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")

        padding = str(int(self.style.table.cell_padding_pt * 20))  # Convert to twips

        for margin_name in ["top", "left", "bottom", "right"]:
            margin = OxmlElement(f"w:{margin_name}")
            margin.set(qn("w:w"), padding)
            margin.set(qn("w:type"), "dxa")
            tcMar.append(margin)

        tcPr.append(tcMar)
